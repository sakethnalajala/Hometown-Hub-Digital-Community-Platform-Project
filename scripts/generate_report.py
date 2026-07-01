# -*- coding: utf-8 -*-
"""
Generates a professional internship project report (.docx) for
"Hometown Hub - Digital Community Platform".
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
STEEL = RGBColor(0x2E, 0x5C, 0x8A)
GREY = RGBColor(0x44, 0x44, 0x44)

doc = Document()

# ---------------------------------------------------------------------------
# Base styles
# ---------------------------------------------------------------------------
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11.5)
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.3

for lvl, size, color in [(1, 16, NAVY), (2, 13, STEEL), (3, 12, STEEL)]:
    st = doc.styles[f'Heading {lvl}']
    st.font.name = 'Calibri'
    st.font.size = Pt(size)
    st.font.color.rgb = color
    st.font.bold = True
    st.paragraph_format.space_before = Pt(14)
    st.paragraph_format.space_after = Pt(6)
    st.paragraph_format.keep_with_next = True

# Page margins
for section in doc.sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.1)
    section.right_margin = Inches(1.0)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def add_page_number_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.text = ''
    run = p.add_run('Page ')
    run.font.size = Pt(9)
    run.font.color.rgb = GREY
    # PAGE field
    fld_begin = OxmlElement('w:fldChar'); fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = 'PAGE'
    fld_end = OxmlElement('w:fldChar'); fld_end.set(qn('w:fldCharType'), 'end')
    r = run._r
    r2 = p.add_run()._r
    r2.append(fld_begin); r2.append(instr); r2.append(fld_end)


def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def body(text, italic=False, bold=False, align=None, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if align:
        p.alignment = align
    return p


def bullets(items, style='List Bullet'):
    for it in items:
        p = doc.add_paragraph(style=style)
        if isinstance(it, tuple):
            run = p.add_run(it[0])
            run.bold = True
            p.add_run(it[1])
        else:
            p.add_run(it)


def page_break():
    doc.add_page_break()


def spacer(n=1):
    for _ in range(n):
        doc.add_paragraph()


# ---------------------------------------------------------------------------
# 1. COVER PAGE
# ---------------------------------------------------------------------------
spacer(2)
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('PROJECT REPORT'); r.bold = True; r.font.size = Pt(18); r.font.color.rgb = STEEL
t.paragraph_format.space_after = Pt(4)

t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('on'); r.italic = True; r.font.size = Pt(13); r.font.color.rgb = GREY

spacer(1)
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('HOMETOWN HUB'); r.bold = True; r.font.size = Pt(30); r.font.color.rgb = NAVY
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('Digital Community Platform'); r.bold = True; r.font.size = Pt(18); r.font.color.rgb = STEEL

spacer(1)
# Decorative rule
hr = doc.add_paragraph(); hr.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = hr.add_run('— ' * 18); r.font.color.rgb = STEEL

spacer(1)
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('A report submitted in partial fulfilment of the requirements of the\nInternship Programme'); r.font.size = Pt(12); r.italic = True; r.font.color.rgb = GREY

spacer(3)
for label, value in [
    ('Submitted By', 'N. Saketh'),
    ('Internship Organisation', 'Unified Mentor'),
    ('Project Title', 'Hometown Hub – Digital Community Platform'),
    ('Submission Date', 'June 2026'),
]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'{label}: '); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = NAVY
    r2 = p.add_run(value); r2.font.size = Pt(13); r2.font.color.rgb = GREY
    p.paragraph_format.space_after = Pt(4)

page_break()

# ---------------------------------------------------------------------------
# 2. CERTIFICATE
# ---------------------------------------------------------------------------
heading('Certificate', 1)
spacer(1)
body('This is to certify that the project report entitled "Hometown Hub – Digital '
     'Community Platform" is a bonafide record of the work carried out by '
     'N. Saketh during the internship programme conducted by Unified Mentor.', )
body('The work presented in this report has been completed by the candidate under '
     'due guidance and supervision, and embodies the candidate\'s own efforts. The '
     'project has been developed in accordance with the objectives laid out at the '
     'commencement of the internship and demonstrates the practical application of '
     'modern front-end web development concepts.')
body('To the best of my knowledge, this report or any part of it has not been '
     'submitted earlier for the award of any other certificate, diploma or degree.')
spacer(3)
tbl = doc.add_table(rows=1, cols=2)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
c = tbl.rows[0].cells
c[0].text = ''
c[1].text = ''
sign = [('_______________________', '_______________________'),
        ('Internship Mentor', 'Candidate Signature'),
        ('Unified Mentor', 'N. Saketh')]
for left, right in sign:
    row = tbl.add_row().cells
    pl = row[0].paragraphs[0]; rl = pl.add_run(left); rl.bold = (left[0] != '_')
    pr = row[1].paragraphs[0]; rr = pr.add_run(right); rr.bold = (right[0] != '_')
page_break()

# ---------------------------------------------------------------------------
# 3. ACKNOWLEDGEMENT
# ---------------------------------------------------------------------------
heading('Acknowledgement', 1)
spacer(1)
body('The completion of this project has been a rewarding learning experience, and '
     'I would like to express my sincere gratitude to all those who supported and '
     'guided me throughout this internship.')
body('First and foremost, I extend my heartfelt thanks to Unified Mentor for '
     'providing me with the opportunity to undertake this internship and to work on '
     'a meaningful, real-world project. The structured guidance and the freedom to '
     'explore ideas helped me grow both technically and professionally.')
body('I am deeply grateful to my mentor for their constant encouragement, valuable '
     'feedback and technical insights, which played a crucial role in shaping the '
     'direction and quality of this project. Their patience and willingness to '
     'clarify concepts greatly enriched my understanding of front-end development.')
body('I would also like to thank my faculty members, family and friends for their '
     'continuous motivation and support during the course of this work. Finally, I '
     'acknowledge the wealth of open documentation, tutorials and community '
     'resources that made the learning process smoother and more enjoyable.')
spacer(2)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run('N. Saketh'); r.bold = True; r.font.color.rgb = NAVY
page_break()

# ---------------------------------------------------------------------------
# 4. ABSTRACT
# ---------------------------------------------------------------------------
heading('Abstract', 1)
spacer(1)
body('Hometown Hub – Digital Community Platform is a web-based application designed '
     'to bring the residents of a town or city closer together through a single, '
     'unified digital space. In an age where civic information, local commerce, '
     'community events and essential public services are scattered across numerous '
     'disconnected platforms, Hometown Hub consolidates these everyday needs into '
     'one accessible, easy-to-use website.')
body('The platform enables users to discover and join local communities, browse and '
     'register for events, buy and sell items through a community marketplace, '
     'access curated educational resources, and locate important government service '
     'information. A clean, modern and fully responsive user interface ensures that '
     'the platform works seamlessly across desktops, tablets and mobile devices.')
body('This project was developed as part of the internship programme at Unified '
     'Mentor with an emphasis on core front-end web technologies. The application '
     'is built using HTML5 for structure, CSS3 and Bootstrap for responsive styling, '
     'and JavaScript (ES6) for interactivity. Browser Local Storage is used to '
     'persist user data such as authentication state, listings and preferences on '
     'the client side, allowing the application to function without a dedicated '
     'backend server.')
body('This report documents the motivation, objectives, system design, modules, '
     'features, implementation details and challenges encountered during the '
     'development of Hometown Hub, and concludes with a discussion of potential '
     'future enhancements.')
page_break()

# ---------------------------------------------------------------------------
# 5. TABLE OF CONTENTS
# ---------------------------------------------------------------------------
heading('Table of Contents', 1)
spacer(1)
toc_p = doc.add_paragraph()
run = toc_p.add_run()
fb = OxmlElement('w:fldChar'); fb.set(qn('w:fldCharType'), 'begin')
instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve')
instr.text = 'TOC \\o "1-3" \\h \\z \\u'
fsep = OxmlElement('w:fldChar'); fsep.set(qn('w:fldCharType'), 'separate')
placeholder = OxmlElement('w:t')
placeholder.text = "Right-click here and choose 'Update Field' to build the Table of Contents."
fend = OxmlElement('w:fldChar'); fend.set(qn('w:fldCharType'), 'end')
run._r.append(fb); run._r.append(instr); run._r.append(fsep); run._r.append(placeholder); run._r.append(fend)
spacer(1)
note = doc.add_paragraph()
nr = note.add_run('Note: The Table of Contents is a live field. On opening the document, '
                  'select the entries above, then press F9 (or right-click → Update Field → '
                  'Update entire table) to populate page numbers automatically.')
nr.italic = True; nr.font.size = Pt(9.5); nr.font.color.rgb = GREY
page_break()

# ---------------------------------------------------------------------------
# 6. INTRODUCTION
# ---------------------------------------------------------------------------
heading('1. Introduction', 1)
body('The way people interact with their local communities has changed dramatically '
     'over the past decade. While global social networks have made it easier to '
     'connect with people across the world, they have, somewhat ironically, made it '
     'harder to stay engaged with what is happening right next door. Information '
     'about a neighbourhood event, a nearby second-hand item for sale, a new '
     'educational programme, or the procedure to apply for a government document is '
     'often fragmented across many different websites, social media groups, '
     'notice-boards and word of mouth.')
body('Hometown Hub – Digital Community Platform was conceived to solve exactly this '
     'problem. It is a centralised, web-based platform that gathers the most common '
     'local needs of a citizen into a single, well-organised destination. Rather '
     'than visiting several disconnected services, a resident can open Hometown Hub '
     'and find communities to join, events to attend, products to buy or sell, '
     'learning opportunities to explore, and government services to access — all '
     'from one place.')
heading('1.1 What is Hometown Hub?', 2)
body('Hometown Hub is a responsive, front-end web application that acts as a digital '
     'town square. It is organised into clearly defined portals — Communities, '
     'Events, Marketplace, Education and Government Services — each catering to a '
     'specific aspect of local life. The platform presents information through '
     'attractive, card-based layouts with images, concise descriptions and clear '
     'call-to-action buttons, making it intuitive for users of all ages and '
     'technical backgrounds.')
heading('1.2 Why was it Developed?', 2)
body('The platform was developed to address the growing disconnect between citizens '
     'and locally relevant information. The key motivations behind the project were:')
bullets([
    ('Centralisation: ', 'to bring scattered local information into one unified, dependable platform.'),
    ('Accessibility: ', 'to provide a clean, responsive interface usable on any device, anywhere.'),
    ('Community engagement: ', 'to encourage people to participate in local events and groups.'),
    ('Local economy: ', 'to support residents in buying and selling within their own community.'),
    ('Awareness: ', 'to make essential education and government service information easy to find.'),
])
heading('1.3 Problems it Solves', 2)
body('By unifying these services, Hometown Hub solves several everyday problems: it '
     'eliminates the need to search across multiple sources, reduces the chance of '
     'missing local events, provides a trusted space for community commerce, and '
     'lowers the barrier to accessing public service information. The result is a '
     'more informed, more connected and more engaged local community.')
page_break()

# ---------------------------------------------------------------------------
# 7. PROBLEM STATEMENT
# ---------------------------------------------------------------------------
heading('2. Problem Statement', 1)
body('In most towns and cities, information that is essential to daily life is '
     'distributed across a wide range of disconnected channels. Community updates '
     'may live in private messaging groups, event details on social media pages, '
     'second-hand goods on classified websites, educational opportunities on '
     'institutional portals, and government procedures on separate official sites. '
     'This fragmentation creates real friction for citizens.')
body('The core problems this project aims to address are as follows:')
bullets([
    ('Information overload and fragmentation: ', 'users must navigate many platforms to gather basic local information.'),
    ('Low community participation: ', 'events and groups are often poorly publicised, leading to low turnout and engagement.'),
    ('Lack of a trusted local marketplace: ', 'buying and selling within a community is informal and difficult to track.'),
    ('Poor discoverability of services: ', 'educational and government services are hard to locate and understand.'),
    ('Inconsistent user experience: ', 'many local platforms are outdated and not optimised for mobile devices.'),
])
body('There is therefore a clear need for a single, modern, responsive platform that '
     'consolidates community, commerce, events, education and civic services into a '
     'coherent and user-friendly experience. Hometown Hub is the proposed solution '
     'to this problem.')
page_break()

# ---------------------------------------------------------------------------
# 8. OBJECTIVES
# ---------------------------------------------------------------------------
heading('3. Objectives', 1)
body('The primary objective of the project is to design and develop a unified '
     'digital community platform that addresses the problems described above. The '
     'specific objectives are:')
bullets([
    ('Connect people with their hometown communities — ', 'enable users to discover, join and engage with local community groups.'),
    ('Provide a marketplace for buying and selling — ', 'offer a simple, trustworthy space for community commerce.'),
    ('Display local events — ', 'present upcoming events with images, dates, locations and registration options.'),
    ('Offer education portal information — ', 'curate learning resources, courses and providers in one accessible portal.'),
    ('Provide government services information — ', 'list essential civic services with clear descriptions and access points.'),
    ('Support responsive design and a modern UI — ', 'ensure a consistent, attractive experience across all screen sizes.'),
])
body('Beyond these functional goals, the project also aims to demonstrate sound '
     'front-end engineering practices, including clean and semantic markup, '
     'reusable styling, modular JavaScript, and client-side data persistence using '
     'Local Storage.')
page_break()

# ---------------------------------------------------------------------------
# 9. SCOPE
# ---------------------------------------------------------------------------
heading('4. Scope of the Project', 1)
body('The scope of Hometown Hub defines the boundaries of what the platform is '
     'designed to deliver within the duration of the internship. The project '
     'focuses on building a complete, functional front-end application with '
     'client-side data handling.')
heading('4.1 In Scope', 2)
bullets([
    'A responsive multi-page web application accessible on desktop, tablet and mobile.',
    'Dedicated portals for Communities, Events, Marketplace, Education and Government Services.',
    'User authentication (login and signup) managed on the client side.',
    'Creation, listing and viewing of community content, events and marketplace items.',
    'Persistent storage of user and content data through browser Local Storage.',
    'A consistent visual identity using a modern component-based card layout.',
])
heading('4.2 Out of Scope', 2)
bullets([
    'A dedicated server-side backend and relational database (proposed as a future enhancement).',
    'Real-time features such as live chat and push notifications.',
    'Online payment processing and order fulfilment.',
    'Native mobile applications for Android and iOS.',
])
body('By clearly defining this scope, the project remains focused and achievable '
     'while still delivering a complete and demonstrable platform that fulfils all '
     'stated objectives.')
page_break()

# ---------------------------------------------------------------------------
# 10. TECHNOLOGIES USED
# ---------------------------------------------------------------------------
heading('5. Technologies Used', 1)
body('Hometown Hub was built primarily with core front-end web technologies, '
     'supported by widely-adopted libraries and design resources. The technology '
     'stack was chosen for its reliability, broad browser support and suitability '
     'for building responsive, interactive interfaces.')

tech_table = doc.add_table(rows=1, cols=2)
tech_table.style = 'Light Grid Accent 1'
tech_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = tech_table.rows[0].cells
hdr[0].paragraphs[0].add_run('Technology').bold = True
hdr[1].paragraphs[0].add_run('Purpose / Role in the Project').bold = True
tech_rows = [
    ('HTML5', 'Provides the semantic structure and content of every page using modern, accessible markup elements.'),
    ('CSS3', 'Handles all visual styling, layout, colours, spacing, transitions and hover animations.'),
    ('JavaScript (ES6)', 'Adds interactivity, dynamic rendering of content, form handling and application logic.'),
    ('Local Storage', 'Persists user sessions, listings and preferences in the browser without a backend server.'),
    ('Bootstrap', 'Supplies a responsive grid system and ready-made UI components for rapid, consistent design.'),
    ('Font Awesome', 'Provides a comprehensive set of scalable vector icons used throughout the interface.'),
    ('Google Fonts', 'Delivers modern, web-optimised typography to enhance readability and visual appeal.'),
]
for name, purpose in tech_rows:
    cells = tech_table.add_row().cells
    cells[0].paragraphs[0].add_run(name).bold = True
    cells[1].paragraphs[0].add_run(purpose)

spacer(1)
body('Together, these technologies enable a lightweight yet powerful application '
     'that runs entirely in the browser, requires no installation, and remains fast '
     'and responsive across devices.')
page_break()

# ---------------------------------------------------------------------------
# 11. SYSTEM ARCHITECTURE
# ---------------------------------------------------------------------------
heading('6. System Architecture', 1)
body('Hometown Hub follows a client-side architecture in which the entire '
     'application logic, presentation and data persistence reside within the user\'s '
     'web browser. This architecture keeps the system simple, portable and fast, '
     'while still supporting rich, interactive functionality.')
heading('6.1 Architectural Layers', 2)
bullets([
    ('Presentation Layer (HTML5 + CSS3 + Bootstrap): ', 'defines the structure and visual appearance of all pages, ensuring a responsive and consistent look.'),
    ('Logic Layer (JavaScript ES6): ', 'controls user interaction, validates forms, renders content dynamically and manages navigation between modules.'),
    ('Data Layer (Local Storage): ', 'stores and retrieves application data such as registered users, community posts, event registrations and marketplace listings.'),
])
heading('6.2 Data Flow', 2)
body('When a user interacts with the application — for example, by registering an '
     'account or creating a marketplace listing — the JavaScript logic layer '
     'captures the input, validates it, and writes the resulting data to Local '
     'Storage as serialised JSON. When a page loads, the same logic layer reads the '
     'stored data, parses it, and dynamically generates the corresponding HTML '
     'elements. This read-render-update cycle allows the platform to behave like a '
     'data-driven application without a traditional server.')
heading('6.3 Conceptual Architecture Diagram', 2)
arch = doc.add_paragraph()
arch.alignment = WD_ALIGN_PARAGRAPH.CENTER
ar = arch.add_run(
    'User (Browser)\n'
    '          ↓  ↑\n'
    'Presentation Layer  —  HTML5 / CSS3 / Bootstrap / Font Awesome\n'
    '          ↓  ↑\n'
    'Logic Layer  —  JavaScript (ES6) Event Handling & Rendering\n'
    '          ↓  ↑\n'
    'Data Layer  —  Browser Local Storage (JSON)')
ar.font.name = 'Consolas'; ar.font.size = Pt(10.5); ar.font.color.rgb = NAVY
body('This layered separation of concerns makes the codebase easier to maintain and '
     'extend, and provides a clear migration path towards a full client-server '
     'architecture in the future.')
page_break()

# ---------------------------------------------------------------------------
# 12. MODULES DESCRIPTION
# ---------------------------------------------------------------------------
heading('7. Modules Description', 1)
body('The application is organised into a number of well-defined modules, each '
     'responsible for a distinct area of functionality. This modular structure '
     'improves clarity, reusability and maintainability.')
modules = [
    ('7.1 Home', 'The Home module is the landing page and central hub of the platform. It '
     'introduces the purpose of Hometown Hub, highlights the available portals through '
     'attractive cards, and provides quick navigation to every other module. It is '
     'designed to make a strong first impression and guide users effortlessly to the '
     'section they need.'),
    ('7.2 Login & Signup', 'This module manages user authentication. New users can create '
     'an account through the signup form, while returning users can log in with their '
     'credentials. Input validation ensures that data is entered correctly, and user '
     'information is stored securely in Local Storage to maintain the session across '
     'pages.'),
    ('7.3 Communities', 'The Communities module allows users to discover and join local '
     'community groups, each representing a different place or interest. Every community '
     'is presented as a card containing an image, a name, a short description and a Join '
     'Community button, encouraging participation and engagement.'),
    ('7.4 Events', 'The Events module displays upcoming local events. Each event card '
     'includes an image, the event date, the location, a brief description and a Register '
     'button. Users can quickly browse what is happening around them and register their '
     'interest with a single click.'),
    ('7.5 Marketplace', 'The Marketplace module provides a space for community commerce. '
     'Users can browse products, each shown with an image, price, seller name and a Buy '
     'Now button. The module supports listing new items for sale and completing a '
     'simulated purchase, fostering a local buy-and-sell economy.'),
    ('7.6 Government Services', 'This module lists essential government services such as '
     'identity, transport and civic services. Each service is presented with an icon, a '
     'clear description and an Access Service button, helping citizens quickly understand '
     'and reach the services they require.'),
    ('7.7 Education Portal', 'The Education module curates learning resources and courses. '
     'Each resource is displayed with a thumbnail image, a description, the provider name '
     'and an Open Course button, making it easy for users to find relevant educational '
     'opportunities.'),
    ('7.8 User Profile', 'The User Profile module presents the logged-in user\'s '
     'information and activity. It serves as a personal dashboard where users can view '
     'their details, manage their account and review their interactions with the '
     'platform.'),
]
for title, text in modules:
    heading(title, 2)
    body(text)
page_break()

# ---------------------------------------------------------------------------
# 13. FEATURES
# ---------------------------------------------------------------------------
heading('8. Features', 1)
body('Hometown Hub offers a comprehensive set of features designed to deliver a '
     'smooth and engaging user experience. The most significant features are '
     'described in detail below.')
features = [
    ('8.1 Responsive Design', 'The platform adapts fluidly to any screen size, from large '
     'desktop monitors to small mobile phones. Using the Bootstrap grid system and CSS3 '
     'media queries, layouts reflow gracefully, ensuring that content remains readable '
     'and controls remain easy to use on every device.'),
    ('8.2 Modern Card-Based UI', 'Content across all portals is presented through visually '
     'consistent cards, each featuring an attractive image, a unique title, a concise '
     'description and a clearly labelled action button. Hover animations provide subtle, '
     'satisfying feedback that enhances interactivity.'),
    ('8.3 User Authentication', 'Secure client-side login and signup functionality allows '
     'users to create accounts and maintain personalised sessions. Form validation '
     'prevents invalid input, and session state is preserved through Local Storage.'),
    ('8.4 Community Engagement', 'Users can browse a curated set of communities, each '
     'representing a different city or interest, and join them with a single click, '
     'fostering a sense of belonging and local participation.'),
    ('8.5 Event Discovery and Registration', 'Upcoming events are showcased with rich '
     'detail, and users can register instantly. This encourages attendance and keeps the '
     'community informed about local happenings.'),
    ('8.6 Community Marketplace', 'A built-in marketplace lets users browse products and '
     'complete purchases smoothly. Each listing clearly displays the price and seller, '
     'and the Buy Now action confirms the purchase without any disruptive errors.'),
    ('8.7 Information Portals', 'The Education and Government Services portals make '
     'essential information easy to discover, with descriptive cards and direct access '
     'buttons. Where an online service is not yet available, a clear, professional '
     'notification informs the user instead of failing silently.'),
    ('8.8 Data Persistence', 'Through Local Storage, user data and content persist between '
     'sessions on the same device, giving the application the feel of a stateful, '
     'data-driven system without requiring a backend server.'),
]
for title, text in features:
    heading(title, 2)
    body(text)
page_break()

# ---------------------------------------------------------------------------
# 14. IMPLEMENTATION
# ---------------------------------------------------------------------------
heading('9. Implementation', 1)
body('The implementation of Hometown Hub brings together the chosen technologies to '
     'create a cohesive, interactive platform. This section explains how each core '
     'technology was applied during development.')
heading('9.1 Structuring with HTML5', 2)
body('HTML5 was used to lay out the semantic structure of every page. Semantic '
     'elements such as headers, navigation bars, sections, articles and footers were '
     'used to organise content meaningfully, improving both accessibility and '
     'maintainability. Each portal — Home, Communities, Events, Marketplace, '
     'Education and Government Services — was structured with consistent, reusable '
     'markup patterns.')
heading('9.2 Styling with CSS3 and Bootstrap', 2)
body('CSS3 governed the visual presentation of the platform, including colours, '
     'typography, spacing, rounded cards, gradients, shadows and smooth hover '
     'animations. The Bootstrap framework provided a responsive 12-column grid and '
     'pre-built components, which accelerated development and guaranteed a '
     'consistent, mobile-friendly layout. Media queries were used to fine-tune the '
     'experience for different breakpoints, while Google Fonts and Font Awesome '
     'icons enhanced the overall aesthetic.')
heading('9.3 Interactivity with JavaScript (ES6)', 2)
body('JavaScript (ES6) powered the dynamic behaviour of the application. Modern '
     'features such as arrow functions, template literals, the let and const '
     'keywords, array methods and JSON handling were used to write clean, concise '
     'logic. JavaScript was responsible for handling form submissions, validating '
     'user input, dynamically generating cards from stored data, managing navigation '
     'and providing real-time feedback through notifications.')
heading('9.4 Data Persistence with Local Storage', 2)
body('Local Storage served as the client-side database for the platform. When a '
     'user signed up, created a listing or registered for an event, the relevant '
     'data was serialised into JSON and saved to Local Storage using the '
     'setItem method. On subsequent visits, the data was retrieved with getItem, '
     'parsed back into JavaScript objects, and rendered onto the page. This approach '
     'allowed the application to remember users and their content across sessions, '
     'simulating the behaviour of a server-backed system entirely within the '
     'browser.')
heading('9.5 Bringing it Together', 2)
body('The combination of these technologies produced a fast, self-contained '
     'application. HTML provided the skeleton, CSS and Bootstrap delivered the '
     'visual polish and responsiveness, JavaScript supplied the intelligence and '
     'interactivity, and Local Storage gave the platform a memory. The result is a '
     'fully functional digital community platform that runs smoothly without any '
     'external dependencies.')
page_break()

# ---------------------------------------------------------------------------
# 15. SCREENSHOTS
# ---------------------------------------------------------------------------
heading('10. Screenshots', 1)
body('This section presents screenshots of the key pages of the Hometown Hub '
     'platform. Each placeholder below should be replaced with the corresponding '
     'screenshot captured from the running application.')
shots = ['Home Page', 'Login Page', 'Marketplace', 'Communities',
         'Events', 'Government Services', 'Education Portal', 'User Profile']
for i, caption in enumerate(shots, 1):
    box = doc.add_paragraph()
    box.alignment = WD_ALIGN_PARAGRAPH.CENTER
    br = box.add_run('\n[  Insert Screenshot Here  ]\n')
    br.italic = True; br.font.size = Pt(12); br.font.color.rgb = GREY
    # simple bordered look via shading paragraph
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(f'Figure 10.{i}: {caption}')
    cr.bold = True; cr.font.size = Pt(10.5); cr.font.color.rgb = NAVY
    spacer(1)
    if i == 4:
        page_break()
page_break()

# ---------------------------------------------------------------------------
# 16. CHALLENGES FACED
# ---------------------------------------------------------------------------
heading('11. Challenges Faced', 1)
body('As with any development project, building Hometown Hub presented a number of '
     'challenges. Overcoming these obstacles was an important part of the learning '
     'experience. The most notable challenges and their solutions are described '
     'below.')
challenges = [
    ('11.1 Achieving True Responsiveness', 'Ensuring that every page looked and worked '
     'well across a wide range of devices was demanding. This was resolved by adopting '
     'the Bootstrap grid system, applying CSS3 media queries at key breakpoints, and '
     'testing the layout repeatedly on different screen sizes.'),
    ('11.2 Managing Data Without a Backend', 'Persisting user accounts and content without '
     'a server required careful design. The challenge was addressed by structuring all '
     'data as JSON and using Local Storage as a lightweight client-side store, with '
     'helper functions to read, write and update records consistently.'),
    ('11.3 Maintaining Visual Consistency', 'Keeping a uniform look across many pages and '
     'components was difficult as the project grew. A shared set of CSS classes and a '
     'reusable card pattern were created to ensure consistent spacing, colours and '
     'typography throughout the platform.'),
    ('11.4 Handling User Input and Validation', 'Preventing invalid or incomplete data '
     'from entering the system was essential. JavaScript form validation was implemented '
     'to check inputs before saving, and clear feedback messages were shown to guide the '
     'user.'),
    ('11.5 Avoiding Broken Images and Errors', 'External images occasionally failed to '
     'load, and certain actions risked producing console errors. These issues were '
     'mitigated by providing fallback images, validating data before rendering, and '
     'replacing silent failures with clear, professional notifications.'),
    ('11.6 Cross-Browser Compatibility', 'Subtle differences between browsers affected '
     'styling and behaviour. Standards-compliant HTML5, CSS3 and ES6 were used, and the '
     'application was tested across multiple browsers to ensure a consistent experience.'),
]
for title, text in challenges:
    heading(title, 2)
    body(text)
page_break()

# ---------------------------------------------------------------------------
# 17. FUTURE ENHANCEMENTS
# ---------------------------------------------------------------------------
heading('12. Future Enhancements', 1)
body('While Hometown Hub already delivers a complete and functional experience, '
     'there is significant scope to extend and enrich the platform in the future. '
     'The following enhancements are envisioned:')
bullets([
    ('Backend Integration: ', 'introduce a server-side backend and database to securely store and manage data at scale.'),
    ('Real-Time Chat: ', 'enable instant messaging between community members for richer interaction.'),
    ('Payment Gateway: ', 'integrate secure online payments to support genuine marketplace transactions.'),
    ('Notifications: ', 'add email and push notifications to keep users informed about events and activity.'),
    ('Maps Integration: ', 'embed interactive maps to display event venues, services and community locations.'),
    ('AI-Based Recommendations: ', 'suggest relevant communities, events and products based on user interests.'),
    ('Mobile Application Support: ', 'develop native Android and iOS applications for an enhanced mobile experience.'),
])
body('These enhancements would transform Hometown Hub from a robust front-end '
     'prototype into a fully-fledged, production-ready digital community platform '
     'capable of serving real users at scale.')
page_break()

# ---------------------------------------------------------------------------
# 18. CONCLUSION
# ---------------------------------------------------------------------------
heading('13. Conclusion', 1)
body('Hometown Hub – Digital Community Platform successfully demonstrates how core '
     'front-end web technologies can be combined to build a meaningful, real-world '
     'application. By consolidating communities, events, a marketplace, education '
     'and government services into a single responsive platform, the project '
     'addresses a genuine need for centralised, locally-relevant information.')
body('Throughout the development of this project, valuable practical experience was '
     'gained in structuring content with HTML5, styling responsive interfaces with '
     'CSS3 and Bootstrap, adding interactivity with JavaScript (ES6), and persisting '
     'data with Local Storage. Equally important were the lessons learned in '
     'problem-solving, user-centred design and writing clean, maintainable code.')
body('The platform meets all of its stated objectives and provides a solid '
     'foundation for future growth. With the proposed enhancements, Hometown Hub has '
     'the potential to become a powerful tool for strengthening local communities in '
     'the digital age. This internship project at Unified Mentor has been an '
     'enriching learning journey that has significantly deepened my understanding of '
     'modern web development.')
page_break()

# ---------------------------------------------------------------------------
# 19. REFERENCES
# ---------------------------------------------------------------------------
heading('14. References', 1)
refs = [
    'MDN Web Docs – HTML, CSS and JavaScript Reference. Mozilla Developer Network. https://developer.mozilla.org',
    'W3Schools – Web Development Tutorials. https://www.w3schools.com',
    'Bootstrap Documentation. https://getbootstrap.com/docs',
    'Font Awesome – Icon Library and Documentation. https://fontawesome.com',
    'Google Fonts. https://fonts.google.com',
    'JavaScript (ES6) Specification – ECMAScript Language Specification. https://www.ecma-international.org',
    'Web Storage API (Local Storage) – MDN Web Docs. https://developer.mozilla.org/docs/Web/API/Web_Storage_API',
    'Unified Mentor Internship Programme Resources and Guidelines.',
]
for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph(style='List Number')
    p.add_run(ref)

# ---------------------------------------------------------------------------
# Footer page numbers on all sections
# ---------------------------------------------------------------------------
for section in doc.sections:
    add_page_number_footer(section)

out = r'c:\Users\saket\OneDrive\Desktop\Hometown Hub – Digital Community Platform\Hometown_Hub_Project_Report.docx'
doc.save(out)
print('SAVED:', out)
